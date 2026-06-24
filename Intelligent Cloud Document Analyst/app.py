import base64
import io

from flask import Flask, request, jsonify
from docx import Document

app = Flask(__name__)

# --- Lookup tables: deterministic business logic (not AI) ---

SEVERITY_SCORE = {
    "low": 10,
    "medium": 35,
    "high": 65,
    "critical": 90
}

CONFIDENCE_MULTIPLIER = {
    "low": 0.7,
    "medium": 0.85,
    "high": 1.0
}

# Maps incident_type -> responsible team
INCIDENT_ROUTING = {
    "brute_force": "SOC Tier 1",
    "unauthorized_access": "SOC Tier 2",
    "malware": "Endpoint Security Team",
    "phishing": "Email Security Team",
    "data_exfiltration": "Incident Response (IR) Team",
    "denial_of_service": "Network Operations Team",
    "other": "SOC Tier 1"
}

SLA_MINUTES_BY_SEVERITY = {
    "low": 480,      # 8 hours
    "medium": 120,   # 2 hours
    "high": 30,
    "critical": 10
}


def calculate_risk_score(severity: str, confidence: str) -> int:
    """Deterministic numeric risk score (0-100) from severity + confidence."""
    base = SEVERITY_SCORE.get(severity.lower(), 10)
    multiplier = CONFIDENCE_MULTIPLIER.get(confidence.lower(), 0.7)
    score = round(base * multiplier)
    return min(max(score, 0), 100)


def get_routing_team(incident_type: str) -> str:
    return INCIDENT_ROUTING.get(incident_type.lower(), "SOC Tier 1")


def get_sla_minutes(severity: str) -> int:
    return SLA_MINUTES_BY_SEVERITY.get(severity.lower(), 480)


@app.route("/extract-docx", methods=["POST"])
def extract_docx():
    """
    Accepts JSON: { "base64_data": "<base64-encoded .docx content>" }
    Returns JSON: { "text": "<extracted plain text>" }
    """
    payload = request.get_json(force=True)
    base64_data = payload.get("base64_data", "")

    if not base64_data:
        return jsonify({"error": "Missing base64_data field"}), 400

    try:
        file_bytes = base64.b64decode(base64_data)
        doc = Document(io.BytesIO(file_bytes))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also pull text out of any tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        extracted_text = "\n".join(paragraphs)
        return jsonify({"text": extracted_text})

    except Exception as e:
        return jsonify({"error": f"Failed to parse DOCX: {str(e)}"}), 400


@app.route("/enrich", methods=["POST"])
def enrich():
    data = request.get_json(force=True)

    severity = data.get("severity", "low")
    confidence = data.get("confidence", "low")
    incident_type = data.get("incident_type", "other")

    risk_score = calculate_risk_score(severity, confidence)
    routing_team = get_routing_team(incident_type)
    sla_minutes = get_sla_minutes(severity)

    enriched = {
        **data,
        "risk_score": risk_score,
        "routing_team": routing_team,
        "sla_minutes": sla_minutes,
        "requires_escalation": risk_score >= 70
    }

    return jsonify(enriched)


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
